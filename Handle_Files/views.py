from django.shortcuts import render
from django.shortcuts import HttpResponse, HttpResponseRedirect
from django.http import FileResponse
import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.shared import Pt
import re
import shutil
import comtypes.client
import pythoncom
from PyPDF2 import PdfFileReader


def get_num_pages(file_path):
    """
    获取文件总页码
    :param file_path: 文件路径
    :return:
    """
    reader = PdfFileReader(file_path)
    # 不解密可能会报错：PyPDF2.utils.PdfReadError: File has not been decrypted
    if reader.isEncrypted:
        reader.decrypt('')
    page_num = reader.getNumPages()
    return page_num


# 根据参数生成word、pdf文件
def create_file(word_size, column_num, position, name):
    # 打开需要转换的文件
    to_be_read = open(position, "r", encoding='UTF-8')
    content = to_be_read.read()

    # 创建docx文件
    document = Document()

    # 改变页边距
    # changing the page margins
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.1)
        section.bottom_margin = Cm(0.1)
        section.left_margin = Cm(0.1)
        section.right_margin = Cm(0.1)

    # 分栏
    section = document.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(column_num))

    # 修改字体格式
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(word_size)
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    para = document.add_paragraph('')
    paragraph_format = para.paragraph_format
    paragraph_format.space_before = 0  # 上行间距
    paragraph_format.space_after = 0  # 下行间距
    paragraph_format.line_spacing = Pt(word_size + 0.1)  # 行距

    # 将文件内容按照行分割
    content = content.split('\n')
    for line in content:
        # 判断是否为标题
        head = re.match('(#{1,5}) +(.*)', line)
        if head is not None:
            # Head = document.add_heading("", level=len(head.group(1)))
            # run = Head.add_run(head.group(2))
            # run.font.name = u'宋体'
            # run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

            size = word_size + 5 - len(head.group(1))
            Head = document.add_paragraph('')
            Head.paragraph_format.line_spacing = Pt(size + 0.1)
            Head.paragraph_format.space_before = 0
            Head.paragraph_format.space_after = 0
            run = Head.add_run(head.group(2))
            run.font.color.rgb = RGBColor(0, 0, 250)
            run.font.size = Pt(size)

            # 标题之后需要建立新的段落
            para = document.add_paragraph('')
            paragraph_format = para.paragraph_format
            paragraph_format.space_before = 0  # 上行间距
            paragraph_format.space_after = 0  # 下行间距
            paragraph_format.line_spacing = Pt(word_size + 0.1)  # 行距
        else:
            para.add_run(line + '\n')

    # 关闭文件
    to_be_read.close()

    # 保存文件
    document.save(r"./transport_files/out/" + name.split('.')[0] + ".docx")
    document.save(r"./transport_files/download/word/" + name.split('.')[0] + ".docx")

    print(os.getcwd())

    # 转pdf
    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject("Word.Application")
    word.Visible = 0

    word_path = os.getcwd() + r"\\transport_files\\download\\word\\" + name.split('.')[0] + ".docx"
    pdf_path = os.getcwd() + r"\\transport_files\\download\\pdf\\" + name.split('.')[0] + ".pdf"
    newpdf = word.Documents.Open(word_path)
    newpdf.SaveAs(pdf_path, FileFormat=17)
    newpdf.Close()
    pagenum = get_num_pages(pdf_path)
    print(pagenum)
    return pagenum


def handle_file(position, name):
    # 打开需要转换的文件
    to_be_read = open(position, "r", encoding='UTF-8')
    content = to_be_read.read()
    to_be_read.close()

    print(len(content))  # 调试
    content_len = len(content)
    # 定义一些常用变量
    word_size = 10  # 字体大小
    column_num = 2  # 栏目数

    if content_len >= 22500:
        word_size, column_num = 4, 4
    elif content_len >= 15000:
        word_size, column_num = 5, 4
    elif content_len >= 10000:
        word_size, column_num = 7, 3

    # 生成文件
    page_num = create_file(word_size, column_num, position, name)
    while (page_num >= 3) and (word_size != 1):
        del_file('./transport_files/download/word')  # 删除先前下载的临时文件
        del_file('./transport_files/download/pdf')
        word_size -= 0.5
        page_num = create_file(word_size, column_num, position, name)


def del_file(filepath):
    """
    删除某一目录下的所有文件或文件夹
    :param filepath: 路径
    :return:
    """
    del_list = os.listdir(filepath)
    for f in del_list:
        file_path = os.path.join(filepath, f)
        if os.path.isfile(file_path):
            os.remove(file_path)
        elif os.path.isdir(file_path):
            shutil.rmtree(file_path)


def download(request):
    for root, dirs, files in os.walk('./transport_files/download/word'):
        filename = files[0]
    file = open('./transport_files/download/word/' + filename, 'rb')
    response = FileResponse(file)
    response['Content-Type'] = 'application/octet-stream'
    response['Content-Disposition'] = 'attachment;filename="' + filename + '"'
    return response


def download_pdf(request):
    for root, dirs, files in os.walk('./transport_files/download/pdf'):
        filename = files[0]
    file = open('./transport_files/download/pdf/' + filename, 'rb')
    response = FileResponse(file)
    response['Content-Type'] = 'application/octet-stream'
    response['Content-Disposition'] = 'attachment;filename="' + filename + '"'
    return response


# 将请求定位到index.html文件中
def index(request):
    if request.method == 'GET':
        return render(request, 'home.html')
    elif request.method == 'POST':
        del_file('./transport_files/download/word')  # 删除先前下载的临时文件
        del_file('./transport_files/download/pdf')
        content = request.FILES.get("upload", None)
        if not content:
            return HttpResponse("没有上传内容")
        position = './transport_files/in/' + content.name
        # 获取上传文件的文件名，并将其存储到指定位置

        storage = open(position, 'wb+')  # 打开存储文件
        for chunk in content.chunks():  # 分块写入文件
            storage.write(chunk)
        storage.close()  # 写入完成后关闭文件

        handle_file(position, content.name)  # 处理文件

        return render(request, 'download.html')  # 返回客户端信息
    else:
        return HttpResponseRedirect('不支持的请求方法')
